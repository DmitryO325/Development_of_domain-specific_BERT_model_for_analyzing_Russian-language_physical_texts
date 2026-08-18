#!/usr/bin/env python3
"""Собрать отчёт по НИР в .docx (шаблон МИФИ; структура: Смоляр + Седов)."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path.home() / "Downloads" / "Шаблон отчета.docx"
OUT = ROOT / "reports" / "Б23-215_Оберемок_НИР.docx"

TOPIC = (
    "«Разработка доменно-специализированной модели BERT "
    "для анализа русскоязычных физических текстов»"
)
STUDENT = "Оберемок Дмитрий Олегович"
GROUP = "Б23-215"
ADVISOR = "Куратов Андрей Сергеевич"
SEMESTER = "6"
YEAR = "2026"
GITHUB = (
    "https://github.com/DmitryO325/"
    "Development_of_domain-specific_BERT_model_for_analyzing_Russian-language_physical_texts"
)


def set_run_font(run, size_pt: int = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.bold = bold


def add_para(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size: int = 12,
    space_after: int = 6,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size, bold)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    sizes = {1: 14, 2: 13, 3: 12}
    add_para(doc, text, bold=True, size=sizes.get(level, 12), space_after=8)


def add_bullet(doc: Document, text: str) -> None:
    add_para(doc, f"• {text}")


def add_problem_statement(doc: Document) -> None:
    """Раздел 2 — обязательная постановка задачи (по образцу Смоляр)."""
    add_heading(doc, "2 ПОСТАНОВКА ЗАДАЧИ", 1)
    add_para(
        doc,
        "Рассматривается задача построения доменно-специализированной языковой "
        "модели для русскоязычных физических текстов. Источником данных служит "
        "корпус научных публикаций D = {d₁, d₂, …, d_N}, собранный из открытых "
        "журналов и порталов (в первую очередь ufn.ru). Каждый документ описывается "
        "кортежем метаданных и текста:",
    )
    add_para(
        doc,
        "dᵢ = (titleᵢ, textᵢ, authorsᵢ, urlᵢ, sectionᵢ, yᵢ), i = 1, …, N, (1)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_para(
        doc,
        "где textᵢ — основной текст статьи (после извлечения из PDF/HTML); "
        "sectionᵢ — рубрика журнала; yᵢ — внешняя разметка (при наличии), "
        "например код подраздела ГРНТИ 29 «Физика» для задач классификации.",
    )
    add_para(
        doc,
        "Текст документа токенизируется отображением τ и представляется "
        "последовательностью токенов длины не более L (для классического BERT "
        "L = 512):",
    )
    add_para(
        doc,
        "τ(textᵢ) = (w₁, w₂, …, w_{Tᵢ}), Tᵢ ≤ L. (2)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_para(
        doc,
        "Пусть θ — параметры предобученного энкодера (ruSciBERT как базовая модель), "
        "h = Enc_θ(w₁, …, w_T) ∈ ℝ^d — векторное представление последовательности "
        "(например, [CLS]-эмбеддинг). Требуется получить модель ruPhysBERT "
        "путём доменной адаптации на корпусе D.",
    )
    add_para(doc, "Выделяются две связанные подзадачи.", bold=False)

    add_heading(doc, "2.1 Дообучение по задаче MLM", 2)
    add_para(
        doc,
        "На этапе continued pretraining минимизируется функция потерь masked "
        "language modeling: для случайно замаскированных позиций M ⊂ {1, …, T} "
        "модель предсказывает исходные токены:",
    )
    add_para(
        doc,
        "L_MLM(θ, D) = − Σ_{dᵢ∈D} Σ_{j∈Mᵢ} log P_θ(w_j | C_j). (3)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_para(
        doc,
        "Здесь C_j — контекст (немаскированные токены) в позиции j. Цель — снизить "
        "perplexity и улучшить качество представлений на физической терминологии "
        "относительно ruBERT и ruSciBERT.",
    )

    add_heading(doc, "2.2 Классификация и оценка представлений", 2)
    add_para(
        doc,
        "Для размеченного подмножества D' ⊂ D с метками yᵢ ∈ {1, …, K} "
        "(подразделы физики по ГРНТИ или рубрики журнала) строится классификатор "
        "на основе эмбеддинга h:",
    )
    add_para(
        doc,
        "ŷ = argmax_k  softmax(W h + b)_k, (4)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_para(
        doc,
        "Качество сравнивается по macro-F1, accuracy, а также по задаче fill-mask "
        "на физических терминах и подзадачам бенчмарка RuSciBench (фильтр ГРНТИ 29).",
    )

    add_heading(doc, "2.3 Гипотезы и критерии успеха", 2)
    add_para(doc, "Формулируются проверяемые гипотезы:")
    for h in [
        "H1: MLM-дообучение ruSciBERT на корпусе D снижает perplexity по сравнению "
        "с ruSciBERT и ruBERT;",
        "H2: на классификации подразделов физики ruPhysBERT ≥ ruSciBERT > ruBERT "
        "по macro-F1;",
        "H3: введение специальных токенов для формул улучшает fill-mask на "
        "физических терминах (абляция: с токенами / без).",
    ]:
        add_bullet(doc, h)

    add_para(doc, "Дополнительную сложность задаче придают факторы:")
    for factor in [
        "ограниченный объём открытых полнотекстовых статей на русском языке "
        "по сравнению с англоязычными ресурсами;",
        "двухколоночная вёрстка PDF и встроенные формулы, затрудняющие "
        "автоматическое извлечение текста;",
        "необходимость OCR при некорректной встроенной кодировке шрифтов журнала;",
        "ограничение контекста L = 512 токенов для классических BERT-моделей;",
        "разнородность стилей: обзоры, экспериментальные статьи, теоретические работы.",
    ]:
        add_bullet(doc, factor)

    add_para(
        doc,
        "Формально требуется построить отображение Φ_θ: Text → ℝ^d и процедуру "
        "дообучения, минимизирующую L_MLM на D при сравнении с базовыми моделями "
        "и максимизирующую macro-F1 на D' в рамках протокола RuSciBench.",
    )


def build_body(doc: Document) -> None:
    add_heading(doc, "СОДЕРЖАНИЕ", 1)
    for line in [
        "1 Введение .............................................................. 2",
        "2 Постановка задачи ................................................... 3",
        "   2.1 Дообучение по задаче MLM ........................................ 3",
        "   2.2 Классификация и оценка представлений .......................... 4",
        "   2.3 Гипотезы и критерии успеха ...................................... 4",
        "3 Обзор литературы .................................................... 5",
        "4 Сбор корпуса и пилотные эксперименты ............................... 6",
        "   4.1 Пилотные эксперименты с ruBERT и ruSciBERT .................... 6",
        "   4.2 Парсер УФН и извлечение текста из PDF ......................... 7",
        "5 Заключение .......................................................... 8",
        "Список литературы ..................................................... 9",
    ]:
        add_para(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)

    doc.add_page_break()

    add_heading(doc, "1 ВВЕДЕНИЕ", 1)
    add_para(
        doc,
        "Анализ больших массивов русскоязычных научных текстов востребован "
        "при построении систем поиска, классификации публикаций, рекомендаций "
        "и автоматической обработки рефератов. Для физических журналов и учебных "
        "материалов характерна специализированная терминология, формулы и строгий "
        "стиль изложения, которые плохо покрываются общими языковыми моделями.",
    )
    add_para(
        doc,
        "В настоящее время для русского языка доступны мощные энкодеры "
        "(ruBERT, ruRoBERTa), однако они обучены на новостях, википедии и художественных "
        "текстах. Научная модель ruSciBERT учитывает стилистику статей, но не "
        "заточена под поддомен физики. В результате при zero-shot сравнении доменов "
        "и при fill-mask общие модели подставляют бытовую лексику вместо физических "
        "терминов, что подтверждает необходимость доменной адаптации.",
    )
    add_para(
        doc,
        "Одной из ключевых задач НИР является разработка ruPhysBERT — "
        "доменно-специализированной модели на базе ruSciBERT с continued pretraining "
        "на корпусе русскоязычных физических публикаций и последующей оценкой "
        "на задачах MLM, fill-mask и классификации.",
    )
    add_para(
        doc,
        "В ходе научно-исследовательской работы за 6 семестр сформулирована "
        "постановка задачи (раздел 2), проведён обзор литературы (раздел 3), "
        "выполнены пилотные эксперименты и реализован пайплайн сбора корпуса "
        "(раздел 4). Исходный код и документация размещены в репозитории: "
        f"{GITHUB}.",
    )

    doc.add_page_break()
    add_problem_statement(doc)
    doc.add_page_break()

    add_heading(doc, "3 ОБЗОР ЛИТЕРАТУРЫ", 1)
    add_para(
        doc,
        "В рамках первого этапа НИР были изучены современные трансформерные "
        "энкодеры для русского языка и научных текстов. В качестве общего baseline "
        "рассматривается ruBERT-base (DeepPavlov, ai-forever): модель обучена "
        "на википедии, новостях и книгах, поддерживает задачу masked language modeling "
        "(MLM), но не оптимизирована для классификации и доменной семантики без "
        "дополнительного fine-tuning. Усиленным вариантом является ruRoBERTa-large; "
        "модель RuModernBERT-base рассматривается как опциональный baseline "
        "(мультиязычность и код).",
    )
    add_para(
        doc,
        "Для научных текстов основным кандидатом выбран ruSciBERT (Sber AI, MLSA Lab): "
        "модель RoBERTa (~123M параметров), обученная на корпусе порядка 6,5 ГБ "
        "русскоязычных научных публикаций (eLibrary, arXiv.ru, диссертации). "
        "Лицензия Apache 2.0 допускает использование в учебных и исследовательских "
        "целях. Дополнительно рассмотрены SciRus-tiny/small (MLSA Lab) и "
        "Giga-Embeddings-instruct для задач retrieval.",
    )
    add_para(
        doc,
        "Для оценки качества доменной адаптации планируется использование "
        "бенчмарка RuSciBench (2024): задачи классификации по рубрикам ГРНТИ, "
        "в том числе раздел 29 «Физика», регрессия метаданных и retrieval. "
        "Отдельно отмечена проблема обработки математических формул: стандартные "
        "BPE-токенизаторы плохо сегментируют LaTeX; в перспективе рассматривается "
        "введение специальных токенов или дообучение токенизатора.",
    )

    doc.add_page_break()

    add_heading(doc, "4 СБОР КОРПУСА И ПИЛОТНЫЕ ЭКСПЕРИМЕНТЫ", 1)

    add_heading(doc, "4.1 Пилотные эксперименты с ruBERT и ruSciBERT", 2)
    add_para(
        doc,
        "В каталоге draft/ подготовлены Jupyter-ноутбуки rubert.ipynb и "
        "ruscibert.ipynb с экспериментами на Hugging Face Transformers. "
        "Проверена загрузка весов, задача fill-mask и zero-shot сравнение "
        "вероятностей для пар физических и нефизических фраз.",
    )
    add_para(
        doc,
        "Результаты пилота согласуются с гипотезами раздела 2: ruBERT на zero-shot "
        "не различает домен «физика / биология»; при fill-mask подставляются "
        "бытовые лексемы, а не термины. ruSciBERT демонстрирует более уместные "
        "подстановки в научном контексте и выбран в качестве базы Φ_θ для "
        "continued pretraining. Полноценное обучение по (3) и оценка по (4) "
        "запланированы после завершения сборки корпуса D.",
    )

    add_heading(doc, "4.2 Парсер журнала «Успехи физических наук» и извлечение текста", 2)
    add_para(
        doc,
        "Исходный Colab-парсер ufn.ru перенесён в репозиторий и оформлен "
        "как модульная подсистема src/collect/: базовые типы Document и запись "
        "JSONL (base.py), обход выпусков и страниц статей (ufn.py), RSS-ленты "
        "УФН и «Квантовая электроника» (rss_feed.py). Точка входа — "
        "scripts/scrape.py с режимами pilot, ufn и rss.",
    )
    add_para(
        doc,
        "Каждая запись корпуса содержит: source, url, title, text, authors, "
        "published, section, pdf_url, language, extra. За 6 семестр собран "
        "пилотный корпус (~20 статей, corpus.jsonl) и основная выборка из "
        "100 статей УФН (ufn_100.jsonl) с локальным кэшем PDF в data/raw/pdf/. "
        "Дополнительно подготовлено подмножество с OCR-текстом (ufn_10_pdf.jsonl) "
        "для отладки пайплайна.",
    )
    add_para(
        doc,
        "Важная особенность ufn.ru: HTML-страница статьи содержит преимущественно "
        "введение (порядка 3 тыс. символов); полный текст размещён в PDF "
        "с двухколоночной вёрсткой. Поэтому для обучения языковой модели "
        "используется извлечение из PDF, а не только HTML.",
    )

    add_para(
        doc,
        "Модуль pdf_text.py реализует цепочку: скачивание PDF → попытка "
        "извлечения текста через PyMuPDF → при «битой» кодировке шрифтов журнала "
        "— распознавание Tesseract (rus+eng). Учитывается макет: на первой "
        "странице одноколоночная шапка (до блока PACS/DOI), далее чтение "
        "слева направо по колонкам; на последующих страницах — сначала левая "
        "колонка целиком, затем правая. Скрипт rebuild_from_pdf.py позволяет "
        "пересобрать JSONL по уже скачанным PDF без повторного обхода сайта.",
    )
    add_para(
        doc,
        "Документация пайплайна оформлена в docs/PARSERS.md, docs/ARCHITECTURE.md "
        "и docs/RESOURCES.md. На следующем этапе планируется: пересборка "
        "ufn_100_pdf.jsonl, EDA корпуса D, continued pretraining по (3), "
        "оценка гипотез H1–H3 и расширение источников до N ≈ 1000 документов. "
        "Вычисления — локальный GPU (PyTorch), при необходимости Colab/Kaggle.",
    )

    doc.add_page_break()

    add_heading(doc, "5 ЗАКЛЮЧЕНИЕ", 1)
    add_para(
        doc,
        "В ходе научно-исследовательской работы за 6 семестр рассмотрена задача "
        "доменной адаптации языковой модели для русскоязычных физических текстов. "
        "Сформулирована постановка задачи: построение корпуса D, минимизация "
        "L_MLM при дообучении ruSciBERT и оценка представлений на классификации "
        "и бенчмарке RuSciBench; сформулированы гипотезы H1–H3.",
    )
    add_para(
        doc,
        "По результатам обзора литературы выбраны базовые модели ruBERT и "
        "ruSciBERT; проведены пилотные эксперименты, подтверждающие необходимость "
        "специализации. Реализован пайплайн сбора корпуса: парсер ufn.ru, "
        "скачивание 100 PDF, извлечение текста с OCR и учётом двухколоночной вёрстки.",
    )
    add_para(
        doc,
        "Создана основа для дальнейшего этапа: дообучение ruPhysBERT, проверка "
        "гипотез по perplexity и macro-F1, абляции по объёму данных и обработке "
        "формул. Результаты будут оформлены в виде сравнительных таблиц и, "
        "при достаточном качестве, подготовлены к публикации.",
    )

    doc.add_page_break()

    add_heading(doc, "Список литературы", 1)
    refs = [
        "[1] Gena P., et al. ruSciBERT: A BERT model for Russian scientific texts // "
        "Doklady Mathematics. — 2022.",
        "[2] RuSciBench: benchmark for Russian scientific texts // "
        "Mathematical Notes of the Russian Academy of Sciences, 2024. "
        "https://link.springer.com/article/10.1134/S1064562424602191",
        "[3] Devlin J., et al. BERT: Pre-training of Deep Bidirectional Transformers "
        "for Language Understanding // NAACL, 2019.",
        "[4] Семантический анализ научных текстов: опыт создания корпуса "
        "и построения языковых моделей // КиберЛенинка.",
        "[5] Репозиторий НИР: " + GITHUB,
        "[6] Hugging Face: ai-forever/ruSciBERT, ai-forever/ruBERT-base, "
        "mlsa-iai-msu-lab/ru_sci_bench_mteb.",
    ]
    for ref in refs:
        add_para(doc, ref, space_after=4)


def _set_paragraph_text(p, text: str) -> None:
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0])
    else:
        run = p.add_run(text)
        set_run_font(run)


def fill_title_page(doc: Document) -> None:
    """Заполнить поля титульного листа в скопированном шаблоне."""
    by_prefix = {
        "по научно-исследовательской работе": (
            f"по научно-исследовательской работе за {SEMESTER} семестр"
        ),
        "на тему:": f"на тему:\n{TOPIC}",
        "Выполнил:": f"Выполнил: {STUDENT}",
        "Группа:": f"Группа: {GROUP}",
        "Научный руководитель:": f"Научный руководитель:\n{ADVISOR}",
        "Москва –": f"Москва – {YEAR}",
    }
    for p in doc.paragraphs[:40]:
        t = p.text.strip()
        for prefix, new in by_prefix.items():
            if t.startswith(prefix) or t == prefix.strip():
                _set_paragraph_text(p, new)
                break


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    if TEMPLATE.exists():
        shutil.copy2(TEMPLATE, OUT)
        doc = Document(str(OUT))
        fill_title_page(doc)
        # Удалить подсказки шаблона в конце титула
        hint_markers = (
            "Отчет должен содержать",
            "Шрифт Times New Roman",
            "Содержание;",
            "Введение;",
            "Заключение.",
        )
        to_remove = [p for p in doc.paragraphs if any(m in p.text for m in hint_markers)]
        for p in to_remove:
            el = p._element
            el.getparent().remove(el)
        doc.add_page_break()
        build_body(doc)
    else:
        doc = Document()
        for line in [
            "Министерство науки и высшего образования Российской Федерации",
            "«НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ ЯДЕРНЫЙ УНИВЕРСИТЕТ «МИФИ»",
            "ИНСТИТУТ ЛАЗЕРНЫХ И ПЛАЗМЕННЫХ ТЕХНОЛОГИЙ",
            "Кафедра прикладной математики № 31",
            "ОТЧЕТ",
            f"по научно-исследовательской работе за {SEMESTER} семестр",
            f"на тему:\n{TOPIC}",
            f"Выполнил: {STUDENT}",
            f"Группа: {GROUP}",
            f"Научный руководитель: {ADVISOR}",
            f"Москва – {YEAR}",
        ]:
            add_para(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_page_break()
        build_body(doc)

    doc.save(str(OUT))
    downloads_copy = Path.home() / "Downloads" / OUT.name
    shutil.copy2(OUT, downloads_copy)
    print(f"Saved: {OUT}")
    print(f"Copied: {downloads_copy}")


if __name__ == "__main__":
    main()
